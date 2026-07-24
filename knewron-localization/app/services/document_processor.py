"""IFU document processing — Design ref: `LOCKED_Design_v1.0.md` §4.1;
`Technical_Design_Document.md` §2.1.2. Story 5.2.

**Text is never extracted.** The platform only pulls embedded images (+
their position/hash) out of the DOCX for the image sub-pipeline, and later
re-injects localized images into Lokalise's translated copy at the same
positions. The original DOCX is otherwise sent to Lokalise as-is.
"""
import hashlib
from dataclasses import dataclass
from io import BytesIO
from typing import Any

from docx import Document
from docx.image.image import Image as DocxImage

from app.pipeline.image_pipeline import ExtractedImage


@dataclass
class DocumentManifest:
    page_count_estimate: int
    image_count: int
    images: list[ExtractedImage]


def extract_images(docx_bytes: bytes) -> DocumentManifest:
    """Extract embedded images + their position/hash (Technical_Design §2.1.2
    `DocumentProcessor.process_ifu` — no text parsing).
    """
    document = Document(BytesIO(docx_bytes))
    images: list[ExtractedImage] = []

    for position_index, shape in enumerate(document.inline_shapes):
        try:
            blob = shape.image.blob
            filename = f"image_{position_index:04d}.{shape.image.ext}"
        except (AttributeError, ValueError):
            continue

        image_hash = hashlib.sha256(blob).hexdigest()
        images.append(
            ExtractedImage(
                image_id=_hash_to_uuid(image_hash),
                image_hash=image_hash,
                image_bytes=blob,
                position={"inline_shape_index": position_index},
                filename=filename,
            )
        )

    return DocumentManifest(
        page_count_estimate=_estimate_pages(document),
        image_count=len(images),
        images=images,
    )


def replace_images(translated_docx_bytes: bytes, localized_images: dict[str, bytes]) -> bytes:
    """Replace the original images in Lokalise's translated DOCX with the
    localized ones, matched **by position** (LOCKED §4.1, Arch §6).

    `localized_images` maps `str(inline_shape_index)` -> new PNG/JPEG bytes.
    python-docx has no public "swap embedded image" API; this overwrites the
    `ImagePart` blob in place (a well-known community workaround) so the
    picture's on-page size/position (defined in the run's XML) is unchanged
    and only the pixel content is replaced.
    """
    document = Document(BytesIO(translated_docx_bytes))

    for position_index, shape in enumerate(document.inline_shapes):
        new_bytes = localized_images.get(str(position_index))
        if new_bytes is None:
            continue
        image_part = shape._inline.graphic.graphicData.pic.blipFill.blip  # noqa: SLF001
        rId = image_part.embed
        part = document.part.related_parts[rId]
        part._blob = new_bytes  # noqa: SLF001 — no public replace-image API
        part._image = DocxImage.from_blob(new_bytes)  # noqa: SLF001 — refresh cached Image

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _estimate_pages(document: Document) -> int:
    # DOCX has no reliable page count without a rendering engine; approximate
    # from paragraph count (Requirements §3.1.2 samples averaged ~1 image/page).
    return max(1, len(document.paragraphs) // 40)


def _hash_to_uuid(image_hash: str) -> Any:
    import uuid

    return uuid.UUID(image_hash[:32])
