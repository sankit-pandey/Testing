"""Assembly Engine — Design ref: `LOCKED_Design_v1.0.md` §11 (project
structure lists `services/assembler`); `Technical_Design_Document.md` §2.1.7.
Story 5.2.

For IFU, assembly is "take Lokalise's translated DOCX and re-inject the
localized images at their original positions" (LOCKED §4.1) — implemented in
`document_processor.replace_images`; this module wraps it with the
completeness checks from Requirements §5.9.2.
"""
from docx import Document
from io import BytesIO

from app.services.ai_reviewer import Finding, review_document_completeness
from app.services.document_processor import replace_images


def assemble_ifu(
    translated_docx_bytes: bytes,
    localized_images: dict[str, bytes],
    *,
    expected_image_count: int,
) -> tuple[bytes, list[Finding]]:
    """Replace images, then run completeness checks (Requirements §5.9.2:
    "validate completeness", "compare with source — no missing content").
    """
    assembled = replace_images(translated_docx_bytes, localized_images)

    original = Document(BytesIO(translated_docx_bytes))
    final = Document(BytesIO(assembled))
    findings = review_document_completeness(
        expected_image_count=expected_image_count,
        assembled_image_count=len(final.inline_shapes),
        expected_tables=len(original.tables),
        actual_tables=len(final.tables),
    )
    return assembled, findings
